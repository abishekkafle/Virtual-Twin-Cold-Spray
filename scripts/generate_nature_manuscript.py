"""Build the Nature-style manuscript DOCX and Markdown source.

Design preset resolved from the Documents skill:
    google_docs_default

Rationale: Nature-style manuscript submissions should be austere and editable.
The script applies the preset explicitly (US Letter, 1 in margins, Arial 11 pt
body, black headings, simple tables) and adds only named scientific overrides:
caption style, compact table text, and embedded high-resolution figure previews.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
FIG_DIR = PAPER / "nature_figures"
TABLE_DIR = PAPER / "tables"

DOCX_OUT = PAPER / "manuscript_nature_final.docx"
MD_OUT = PAPER / "manuscript_nature_final.md"

P4_DATASET = ROOT / "database" / "cel_p4_simulation_surrogate_dataset.csv"
P5_METRICS = ROOT / "reports" / "cel_p5_surrogate_metrics.json"
P4_RESULTS = ROOT / "extracted" / "production" / "CEL_P4" / "cel_p4_results.json"
WEBXR_VERIFICATION = ROOT / "reports" / "cel_p5_webxr_runtime_verification.json"
KG_WEBXR_MANIFEST = ROOT / "webxr" / "cel_p5_kg_webxr_manifest.json"
MODEL_CARD = ROOT / "models" / "cel_p5_surrogate_model_card.json"
MATERIAL_CROSSWALK = ROOT / "config" / "material_kg_crosswalk.csv"


TITLE = "A knowledge-grounded virtual twin for bounded cold-spray process simulation"
SHORT_TITLE = "Knowledge-grounded cold-spray virtual twin"
AUTHOR_LINE = "Abishek Kafle1,*"
AFFILIATION_LINE = "1 Department of Mechanical Engineering, University of Houston, Houston, TX, USA. *Correspondence: Abishek Kafle."

BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
BORDER = "DADCE0"
HEADER_FILL = "F4F6F9"


REFERENCES = [
    "Papyrin, A., Kosarev, V., Klinkov, S., Alkhimov, A. & Fomin, V. Cold Spray Technology. Elsevier (2007).",
    "Assadi, H., Gärtner, F., Stoltenhoff, T. & Kreye, H. Bonding mechanism in cold gas spraying. Acta Materialia 51, 4379-4394 (2003). https://doi.org/10.1016/S1359-6454(03)00274-X.",
    "Schmidt, T., Gärtner, F., Assadi, H. & Kreye, H. Development of a generalized parameter window for cold spray deposition. Acta Materialia 54, 729-742 (2006). https://doi.org/10.1016/j.actamat.2005.10.005.",
    "Assadi, H., Kreye, H., Gärtner, F. & Klassen, T. Cold spraying - A materials perspective. Acta Materialia 116, 382-407 (2016). https://doi.org/10.1016/j.actamat.2016.06.034.",
    "Johnson, G. R. & Cook, W. H. A constitutive model and data for metals subjected to large strains, high strain rates and high temperatures. Proc. 7th International Symposium on Ballistics, 541-547 (1983).",
    "Kritzinger, W., Karner, M., Traar, G., Henjes, J. & Sihn, W. Digital Twin in manufacturing: A categorical literature review and classification. IFAC-PapersOnLine 51, 1016-1022 (2018). https://doi.org/10.1016/j.ifacol.2018.08.474.",
    "Fuller, A., Fan, Z., Day, C. & Barlow, C. Digital Twin: Enabling technologies, challenges and open research. IEEE Access 8, 108952-108971 (2020). https://doi.org/10.1109/ACCESS.2020.2998358.",
    "Hogan, A. et al. Knowledge graphs. ACM Computing Surveys 54, 71:1-71:37 (2021). https://doi.org/10.1145/3447772.",
    "Geurts, P., Ernst, D. & Wehenkel, L. Extremely randomized trees. Machine Learning 63, 3-42 (2006). https://doi.org/10.1007/s10994-006-6226-1.",
    "Pedregosa, F. et al. Scikit-learn: Machine learning in Python. Journal of Machine Learning Research 12, 2825-2830 (2011).",
    "He, P., Gao, J. & Chen, W. DeBERTaV3: Improving DeBERTa using ELECTRA-style pre-training with gradient-disentangled embedding sharing. ICLR (2023). arXiv:2111.09543.",
    "World Wide Web Consortium. WebXR Device API. Candidate Recommendation Draft. https://www.w3.org/TR/webxr/.",
    "Dassault Systèmes. Abaqus Analysis User's Guide. Dassault Systèmes Simulia Corp. (accessed 2026).",
]


def load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def superscript_citation_runs(paragraph, text: str, style: str | None = None) -> None:
    """Add text to paragraph, converting [[1,2]] tokens to superscript citations."""
    pattern = re.compile(r"\[\[([0-9,\-\s]+)\]\]")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            format_run(run)
        cite = match.group(1).replace(" ", "")
        run = paragraph.add_run(cite)
        format_run(run)
        run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        format_run(run)


def format_run(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def paragraph(text: str = "", style: str | None = None, before: float | None = None, after: float | None = None, align: WD_ALIGN_PARAGRAPH | None = None):
    p = DOC.add_paragraph(style=style)
    if text:
        superscript_citation_runs(p, text)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: Sequence[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def mark_repeating_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table_from_df(df: pd.DataFrame, caption: str, widths_in: Sequence[float], font_size: float = 8.5) -> None:
    cap = paragraph(caption, style="Caption", before=8, after=4)
    cap.runs[0].bold = True
    table = DOC.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    mark_repeating_header(table.rows[0])
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
        set_cell_shading(hdr[j], HEADER_FILL)
        set_cell_border(hdr[j])
        set_cell_margins(hdr[j])
        hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in hdr[j].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                format_run(run, size=font_size, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])
            set_cell_border(cells[j])
            set_cell_margins(cells[j])
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[j].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    format_run(run, size=font_size)
    set_table_width(table, widths_in)
    paragraph("", after=4)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, BLACK),
        ("Heading 2", 16, 18, 6, BLACK),
        ("Heading 3", 14, 16, 4, MUTED),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = False
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    cap = styles["Caption"]
    cap.font.name = "Arial"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    cap.font.size = Pt(9)
    cap.font.color.rgb = BLACK
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.text = f"{SHORT_TITLE} | Page "
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in footer.runs:
        format_run(r, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = "1"
    run.append(txt)
    fld.append(run)
    footer._p.append(fld)


def add_title_page() -> None:
    p = paragraph(TITLE, before=0, after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        format_run(run, size=26, bold=False, color=BLACK)
    paragraph(AUTHOR_LINE, after=3)
    paragraph(AFFILIATION_LINE, after=14)
    paragraph("Manuscript type: Nature Portfolio Article", after=2)
    paragraph("Display items: six main figures, five manuscript/extended-data tables, and source-data files.", after=12)
    note = paragraph(
        "Scope note: this manuscript supports a qualified simulation-surrogate virtual twin. It does not claim experimental physical validation, bond/no-bond thresholds, unseen-pair prediction, or autonomous process control.",
        after=12,
    )
    for run in note.runs:
        format_run(run, italic=True, color=MUTED)


def add_main_text(summary: Dict[str, Any]) -> None:
    p = DOC.add_heading("Summary", level=1)
    paragraph(
        "Engineering virtual twins increasingly connect knowledge bases, high-fidelity simulations, machine-learning inference and immersive interfaces, but these layers are often validated separately. We present a knowledge-grounded simulation-surrogate virtual twin for metal-on-metal cold-spray additive manufacturing. The system links a literature/operator knowledge graph, a gated Abaqus/Explicit coupled Eulerian-Lagrangian simulation campaign, an ExtraTrees surrogate and a WebXR decision-support interface. The verified dataset contains 44 accepted impact simulations over four material-pair domains and eleven velocity levels per pair. All cases passed frozen numerical gates and were retained as machine-learning candidates, with constitutive-review flags preserved as metadata. The selected surrogate achieved pair-aware velocity-interpolation mean R2 = 0.9694 and mean NRMSE = 0.0363 across nine simulation response quantities, whereas leave-one-material-pair-out auditing gave mean R2 = 0.0381. The browser runtime exactly replayed the Python model over 396 case-target comparisons and blocks unsupported material pairs and out-of-range velocities. The result is a bounded, auditable virtual twin for simulation-based decision support rather than an unconstrained physical bonding predictor.",
    )

    DOC.add_heading("Main text", level=1)
    paragraph(
        "Cold spray additive manufacturing deposits solid metal particles by accelerating powders to high velocity and converting impact energy into severe plastic deformation, interfacial heating and, when the process window is favourable, adhesion or buildup.[[1,2,3,4]] These events occupy a regime of large strain, high strain rate, material asymmetry and short time scale, making the process attractive for repair and solid-state additive manufacturing but difficult to generalize from sparse experiments alone.[[4]] Numerical simulation and data-driven modelling can help explore this space, yet a fast predictor becomes scientifically risky when detached from the provenance and applicability boundary of the data that trained it.",
    )
    paragraph(
        "Digital-twin research has emphasized the integration of physical assets, models, data streams and decision support.[[6,7]] In materials processing, however, the boundary between a digital model, a simulation surrogate and a closed-loop digital twin is often blurred. That distinction is consequential for cold spray: a model that interpolates a qualified simulation dataset can be useful for operator training and process exploration, but it does not by itself establish a physical bonding threshold or authorize autonomous machine control. We therefore designed the system around explicit governance: every prediction is tied to a knowledge-graph material identity, a qualified finite-element simulation domain and a runtime gate that can refuse unsupported requests.",
    )

    DOC.add_heading("A claim-bounded virtual-twin architecture", level=2)
    paragraph(
        "The implemented architecture contains five coupled layers: a knowledge graph for materials and operator context, a material-name crosswalk, a gated Abaqus/Explicit CEL simulation layer, a browser-exportable surrogate model and a WebXR human-machine interface (Fig. 1a). The knowledge graph contributes provenance and semantic alignment, not simulation labels. The simulation layer contributes response quantities and numerical-gate metadata. The machine-learning layer contributes fast interpolation only inside qualified material-pair and velocity domains. The runtime and HMI expose those results with explicit warnings rather than hiding model scope inside code.",
    )
    paragraph(
        "This architecture is deliberately asymmetric: allowed outputs are narrow, and blocked claims are broad (Fig. 1b). The virtual twin may display qualified-pair Abaqus/CEL response interpolation, exact browser-executed surrogate inference and KG-backed provenance panels. It may not claim experimental validation, bond/no-bond classification, unseen-pair prediction or autonomous process control. This is not a rhetorical caveat; the same distinction is enforced by the runtime policy described below.",
    )

    DOC.add_heading("A numerically qualified simulation dataset", level=2)
    paragraph(
        "The simulation campaign contains 44 accepted CEL impact cases spanning four metal-on-metal domains: Al6061->SS304, Cu->Cu, Inconel718->Ti6Al4V and Ti6Al4V->Ti6Al4V. Each pair contains 11 velocity levels within the pair-specific qualified range (Table 1 and Fig. 2a). The CEL-P4 production gate extracted all 44 cases, reused 11 previously qualified anchor cases and added 33 new solves. All 44 cases passed the numerical acceptance gate and met the machine-learning candidate criteria.",
    )
    paragraph(
        "The acceptance gates were designed to catch numerical pathologies before any surrogate training. Worst-case final ALLAE/ALLIE was 0.049474 against a threshold of 0.05, absolute total-energy drift normalized by initial kinetic energy was 0.015299 against a threshold of 0.02, particle-volume change was 0.002520 against a threshold of 0.01 and endpoint boundary-return ratio was 0.673428 against a threshold of 0.8 (Fig. 2b and Table 2). The maximum temperature-to-melting-temperature ratio was 0.999966, just below the exclusion threshold. We therefore retained seven constitutive-review cases as flagged data rather than silently discarding them or treating them as unqualified.",
    )

    DOC.add_heading("Response manifolds reveal pair-dependent impact behaviour", level=2)
    paragraph(
        "The qualified CEL responses show material-pair-dependent manifolds rather than a single universal velocity trend (Fig. 3). Terminal particle velocity rises strongly for the high-strength Inconel718->Ti6Al4V and Ti6Al4V->Ti6Al4V cases, remains lower for Cu->Cu and shows non-monotonic behaviour for Al6061->SS304 over its narrower velocity window. Particle flattening, normalized crater depth and maximum homologous temperature likewise separate by pair. These trends motivate a material-pair-aware surrogate and argue against extrapolating a model trained on one pair to another without explicit validation.",
    )

    DOC.add_heading("A fast surrogate for interpolation, not unseen-pair generalization", level=2)
    paragraph(
        "The selected model is an ExtraTrees ensemble trained on categorical material-pair labels, material-registry properties and dimensionless physics features.[[9,10]] The primary validation regime held out velocity levels while preserving representation from every qualified pair in the training folds. Under this intended interpolation test, the selected model achieved mean R2 = 0.9694 and mean NRMSE = 0.0363 across nine response targets. Target-level R2 values were 0.9969 for terminal velocity, 0.9939 for particle flattening, 0.9943 for normalized crater depth, 0.9885 for particle PEEQ p95, 0.9941 for substrate PEEQ p95, 0.9821 for particle maximum temperature, 0.9947 for substrate maximum temperature, 0.9863 for maximum T/Tm and 0.7939 for peak contact pressure (Fig. 4 and Table 4).",
    )
    paragraph(
        "A second validation regime withheld each material pair in turn. This leave-one-pair-out audit is not the intended deployment condition; it is a boundary stress test. The mean R2 fell to 0.0381 for the selected ExtraTrees model, with other candidate models also failing to provide robust unseen-pair prediction (Fig. 5a and Table 3). We therefore treat the surrogate as a qualified-domain interpolator. The model card and runtime gates block unsupported pairs and velocities outside the solved domain instead of returning apparently precise but unvalidated values.",
    )

    DOC.add_heading("WebXR execution with runtime equivalence and traceability", level=2)
    paragraph(
        "The trained tree ensemble was exported to a zero-dependency JavaScript bundle and evaluated inside the WebXR virtual-twin stack.[[12]] Runtime verification replayed the 44-case Python reference table across nine targets, for 396 compared values, with global maximum absolute and relative errors of zero (Fig. 5c). The same runtime returned UNSUPPORTED_PAIR with a null prediction for unsupported material pairs and OUTSIDE_QUALIFIED_VELOCITY_RANGE with a null prediction for out-of-range velocities (Fig. 5b and Table 5).",
    )
    paragraph(
        "A representative Cu->Cu query at 575 m s-1 illustrates the traceability chain (Fig. 6). The HMI request maps the simulation material name to the KG canonical material identifier, resolves the qualified CEL case, records gate metadata and displays the in-domain surrogate response. For that example, the displayed outputs include terminal velocity of 35.88 m s-1, particle flattening of 41.43%, normalized crater depth of 0.452 and maximum T/Tm of 0.575. The panel labels the output as an authorized simulation surrogate, not as a physical deposition decision.",
    )

    DOC.add_heading("Discussion", level=1)
    paragraph(
        "The main contribution is not a new cold-spray bonding law. It is a traceable virtual-twin pattern for using expensive simulations inside an immersive decision-support environment without overstating their authority. The knowledge graph provides evidence structure and material identity; the CEL campaign provides numerically qualified response labels; the surrogate provides subinteractive prediction inside the solved domain; and the WebXR runtime exposes both prediction and applicability state to a human user.",
    )
    paragraph(
        "The results also show why deployment-domain separation matters. Within the qualified domains, the surrogate accurately interpolates most response quantities, especially deformation, velocity and temperature metrics. Peak contact pressure is weaker, consistent with the difficulty of learning a local transient maximum from a small dataset. Across unseen material pairs, however, the leave-one-pair-out audit collapses. A less explicit workflow could easily hide this failure behind a high average interpolation score. Here it is surfaced as a design constraint and implemented as a runtime block.",
    )
    paragraph(
        "Several limitations remain. First, the current results are simulation-surrogate results and require external experimental validation before any physical bonding or process-qualification claim. Second, the dataset contains four qualified material-pair domains and 44 total cases; it is adequate for bounded interpolation studies but not for universal cross-material modelling. Third, seven cases approach near-melt or high-PEEQ constitutive regimes and are retained with explicit review flags. Fourth, the WebXR panel presented here is a decision-support interface and not a closed-loop controller. Future work should add experimentally measured deposit/no-deposit outcomes, extend the solved material-pair matrix and connect the virtual twin to controlled machine telemetry under human-supervised protocols.",
    )


def add_methods() -> None:
    DOC.add_heading("Methods", level=1)
    paragraph(
        "Knowledge-graph and crosswalk layer. The knowledge graph was treated as the provenance and HMI context layer. The packaged manifest records 2,070 merged literature entities, 3,631 NLI-scored literature triples, 1,401 operator entities, 1,544 operator triples and 117 bridge triples. The literature KG was produced by LLM-assisted extraction and subsequent evidence-support scoring using DeBERTa-v3-style NLI verification.[[8,11]] For the virtual twin, a material crosswalk mapped simulation identifiers to KG canonical names for the five materials required by the qualified domains: Al6061, Cu, SS304, Ti6Al4V and Inconel718. The KG was not used to generate simulation labels or to expand the qualified surrogate domain.",
    )
    paragraph(
        "CEL impact simulations. The response dataset was extracted from Abaqus/Explicit coupled Eulerian-Lagrangian impact simulations.[[13]] The constitutive representation used Johnson-Cook-style plasticity with thermal softening and strain-rate sensitivity.[[5]] The production design of experiments contained four metal-on-metal particle/substrate domains, each sampled at eleven velocity levels. The extracted response targets were terminal particle volume-weighted velocity, particle axial flattening, normalized crater depth, particle PEEQ p95, substrate PEEQ p95, particle maximum temperature, substrate maximum temperature, maximum homologous temperature and peak contact pressure.",
    )
    paragraph(
        "Numerical acceptance gates. A case was admitted to the machine-learning candidate set only after passing solver completion and numerical-gate checks. The frozen gate thresholds were final ALLAE/ALLIE <= 0.05, |Delta ETOTAL|/KE0 <= 0.02, particle material volume-change magnitude <= 0.01, endpoint boundary-return ratio <= 0.8 and maximum T/Tm < 1.0 for ML candidacy. Large-strain or near-melt constitutive caution did not automatically exclude a case; instead, flags were preserved as metadata so the surrogate and HMI could expose them.",
    )
    paragraph(
        "Feature engineering and model fitting. The final feature table combined categorical material identifiers, material registry properties and dimensionless physics features. The engineered features were",
    )
    eq = paragraph("Ek* = (0.5 rho_p v^2) / Ap;   H_A = Ap / As;   theta_p = Tref,p / Tm,p;   theta_s = Tref,s / Tm,s.", before=4, after=8)
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in eq.runs:
        format_run(run, size=10, italic=True)
    paragraph(
        "Here rho_p is particle density, v is impact velocity, Ap and As are the particle and substrate Johnson-Cook quasi-static yield parameters, and Tref and Tm are reference and melting temperatures. Candidate regressors were RidgeCV, RandomForest, ExtraTrees and GradientBoosting. ExtraTrees was selected because it provided the highest interpolation mean R2 and could be exported as a deterministic tree ensemble for browser execution.",
    )
    paragraph(
        "Validation regimes. Pair-aware velocity interpolation used five folds formed from velocity_level_index modulo five. Each test fold held out velocity levels while training retained examples from every qualified material pair. Leave-one-pair-out validation withheld each material pair entirely once and was interpreted as a boundary audit for unsupported cross-material prediction, not as the intended deployment condition.",
    )
    paragraph(
        "WebXR runtime verification. The selected ExtraTrees ensemble was serialized to JSON and evaluated by a zero-dependency JavaScript runtime. The runtime was replayed against the Python fitted-model predictions for all 44 cases and nine targets. Domain-gate tests verified that unsupported material pairs and out-of-range velocities return null predictions with explicit status codes rather than extrapolated response values.",
    )
    paragraph(
        "Software and reproducibility. Dataset processing and figure generation used Python 3.12.3, pandas 2.2.2, NumPy 2.2.6, scikit-learn 1.6.1, matplotlib 3.10.8, python-docx 1.2.0 and Pillow 10.3.0. The repository stores extracted CSV/JSON artifacts and reproducibility scripts; Abaqus ODB files are excluded because of size, with extracted histories and gate reports serving as the verification interface.",
    )


def add_availability_and_meta() -> None:
    DOC.add_heading("Data availability", level=1)
    paragraph(
        "The extracted datasets, model metrics, WebXR runtime-verification artifacts, source data for figures and manuscript-generation scripts are available in the public repository: https://github.com/abishekkafle/Virtual-Twin-Cold-Spray. A permanent archival DOI should be minted before journal submission. Raw Abaqus ODB files are not included because of file size; the repository provides extracted case histories, numerical-gate reports and source-data tables needed to reproduce the reported manuscript figures and surrogate metrics.",
    )
    DOC.add_heading("Code availability", level=1)
    paragraph(
        "Code for data packaging, surrogate verification, Nature-style figure generation and manuscript generation is included in the same repository. The browser-executable tree ensemble and WebXR runtime are provided under `webxr/`; manuscript scripts are provided under `scripts/`.",
    )
    DOC.add_heading("Acknowledgements", level=1)
    paragraph("The author thanks the University of Houston research environment and the cold-spray knowledge-graph and virtual-twin development workflow that enabled this manuscript package. Funding acknowledgements should be added before submission.")
    DOC.add_heading("Author contributions", level=1)
    paragraph("A.K. conceived the study, generated and curated the simulation-surrogate dataset, developed the knowledge-graph/WebXR integration, analysed the results and prepared the manuscript. Additional author roles should be added before submission if applicable.")
    DOC.add_heading("Competing interests", level=1)
    paragraph("The author declares no competing interests. This statement should be confirmed before submission.")


def add_references() -> None:
    DOC.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        p = paragraph(f"{i}. {ref}", after=4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)


FIGURES = [
    (
        "Figure 1",
        "Knowledge-grounded virtual-twin architecture.",
        "a, Architecture connecting KG provenance, material crosswalk, qualified CEL simulation, surrogate inference and WebXR decision support. b, Explicit claim boundary separating authorized simulation-surrogate outputs from unsupported physical, unseen-pair and autonomous-control claims.",
        "nature_fig01_virtual_twin_architecture.png",
    ),
    (
        "Figure 2",
        "Numerical qualification of the CEL simulation campaign.",
        "a, Qualified material-pair and velocity domains for the 44-case CEL-P4 dataset. b, Worst-case numerical gate values normalized by their thresholds. c, Maximum homologous temperature with near-melt review flags retained as metadata.",
        "nature_fig02_simulation_qualification.png",
    ),
    (
        "Figure 3",
        "CEL response manifolds over qualified metal-on-metal domains.",
        "a-d, Simulation response trends for terminal velocity, particle flattening, normalized crater depth and maximum homologous temperature across the four qualified material-pair domains.",
        "nature_fig03_cel_response_manifolds.png",
    ),
    (
        "Figure 4",
        "Simulation-surrogate interpolation accuracy.",
        "a, Target-level interpolation R2 values for the selected ExtraTrees surrogate. b-e, Pair-aware velocity-interpolation parity for representative response quantities. The peak-pressure target is retained as a diagnostic because it is visibly less accurate than the deformation, velocity and temperature outputs.",
        "nature_fig04_surrogate_validation.png",
    ),
    (
        "Figure 5",
        "Deployment-domain gating and browser runtime equivalence.",
        "a, Comparison of pair-aware interpolation and leave-one-pair-out boundary auditing across candidate model families. b, Runtime authorization policy. c, JavaScript/Python replay verification over 44 cases and nine targets.",
        "nature_fig05_deployment_gating.png",
    ),
    (
        "Figure 6",
        "WebXR traceability from material selection to displayed prediction.",
        "a, Traceability path from a Cu->Cu, 575 m s-1 query through KG crosswalk, CEL provenance, surrogate inference and HMI output. b, WebXR panel mock-up displaying bounded simulation-surrogate predictions and authorization state.",
        "nature_fig06_webxr_traceability.png",
    ),
]


def add_figure_legends_and_figures() -> None:
    DOC.add_heading("Figure legends", level=1)
    for number, title, legend, filename in FIGURES:
        paragraph(f"{number}. {title} {legend}", style="Caption", before=6, after=6)
    DOC.add_page_break()
    DOC.add_heading("Figures", level=1)
    for number, title, legend, filename in FIGURES:
        path = FIG_DIR / filename
        DOC.add_picture(str(path), width=Inches(6.45))
        inline_shape = DOC.inline_shapes[-1]
        inline_shape._inline.docPr.set("descr", f"{number}. {title} {legend}")
        inline_shape._inline.docPr.set("title", title)
        last = DOC.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph(f"{number}. {title}", style="Caption", before=3, after=2)
        paragraph(legend, style="Caption", before=0, after=8)
        if number != "Figure 6":
            DOC.add_page_break()


def table_dataframes() -> List[Tuple[str, pd.DataFrame, Sequence[float], float]]:
    t1 = pd.read_csv(TABLE_DIR / "table01_qualified_pair_domains.csv")
    t1 = t1[["Pair", "Velocity range (m/s)", "Cases", "Review-flag cases", "Deployment status"]]
    t2 = pd.read_csv(TABLE_DIR / "table02_numerical_acceptance_gates.csv")
    t2 = t2[["Gate", "Threshold", "Observed worst / value", "Status"]]
    t3 = pd.read_csv(TABLE_DIR / "table03_surrogate_model_comparison.csv")
    t4 = pd.read_csv(TABLE_DIR / "table04_selected_model_target_metrics.csv")
    t4 = t4[["Target", "R²", "MAE", "NRMSE", "Observed range"]]
    t5 = pd.read_csv(TABLE_DIR / "table05_webxr_runtime_authorization_policy.csv")
    t5 = t5[["Input condition", "Runtime status", "Output", "Manuscript claim"]]
    return [
        ("Table 1. Qualified material-pair domains used by the simulation surrogate.", t1, [1.65, 1.35, 0.55, 0.95, 2.0], 8.5),
        ("Table 2. Numerical acceptance and ML-candidate gates.", t2, [1.65, 1.55, 1.75, 1.35], 8.2),
        ("Table 3. Model comparison across validation regimes.", t3, [1.45, 1.35, 1.35, 1.15, 1.20], 8.2),
        ("Table 4. Selected ExtraTrees interpolation metrics by response target.", t4, [1.75, 0.55, 0.95, 0.75, 2.4], 7.5),
        ("Table 5. WebXR runtime authorization and deployment policy.", t5, [1.65, 1.75, 1.1, 1.85], 7.6),
    ]


def add_tables() -> None:
    DOC.add_page_break()
    DOC.add_heading("Tables", level=1)
    for caption, df, widths, font_size in table_dataframes():
        add_table_from_df(df, caption, widths, font_size=font_size)


def markdown_source() -> str:
    """A compact markdown source mirroring the DOCX manuscript."""
    df = pd.read_csv(TABLE_DIR / "table01_qualified_pair_domains.csv")
    abstract = (
        "Engineering virtual twins increasingly connect knowledge bases, high-fidelity simulations, machine-learning inference and immersive interfaces, "
        "but these layers are often validated separately. We present a knowledge-grounded simulation-surrogate virtual twin for metal-on-metal cold-spray additive manufacturing. "
        "The verified dataset contains 44 accepted impact simulations over four material-pair domains. The selected ExtraTrees surrogate achieved pair-aware velocity-interpolation "
        "mean R2 = 0.9694 and mean NRMSE = 0.0363 across nine response quantities, while leave-one-material-pair-out auditing gave mean R2 = 0.0381. "
        "The browser runtime exactly replayed the Python model over 396 comparisons and blocks unsupported material pairs and out-of-range velocities."
    )
    lines = [
        f"# {TITLE}",
        "",
        AUTHOR_LINE,
        "",
        AFFILIATION_LINE,
        "",
        "## Summary",
        "",
        abstract,
        "",
        "## Manuscript source note",
        "",
        "The editable DOCX is the authoritative final manuscript package. This Markdown file is a compact source companion generated from the same script.",
        "",
        "## Key quantitative claims",
        "",
        "- 44/44 CEL-P4 cases numerically accepted and ML-candidate.",
        "- ExtraTrees pair-aware interpolation mean R2 = 0.9694; mean NRMSE = 0.0363.",
        "- Leave-one-pair-out mean R2 = 0.0381; unseen-pair prediction is blocked.",
        "- WebXR runtime replay: 44 cases x 9 targets = 396 comparisons; max drift = 0.",
        "",
        "## Figures",
        "",
    ]
    for number, title, legend, filename in FIGURES:
        lines.append(f"- **{number}. {title}** {legend} File: `paper/nature_figures/{filename}`")
    lines.extend(["", "## References", ""])
    for i, ref in enumerate(REFERENCES, 1):
        lines.append(f"{i}. {ref}")
    return "\n".join(lines) + "\n"


def build() -> None:
    global DOC
    DOC = Document()
    style_document(DOC)
    summary = {
        "df": pd.read_csv(P4_DATASET),
        "metrics": load_json(P5_METRICS),
        "p4": load_json(P4_RESULTS),
        "runtime": load_json(WEBXR_VERIFICATION),
        "manifest": load_json(KG_WEBXR_MANIFEST),
        "model_card": load_json(MODEL_CARD),
    }
    add_title_page()
    add_main_text(summary)
    add_methods()
    add_availability_and_meta()
    add_references()
    add_figure_legends_and_figures()
    add_tables()
    DOC.save(DOCX_OUT)
    MD_OUT.write_text(markdown_source(), encoding="utf-8")


if __name__ == "__main__":
    build()
    print(f"Wrote {DOCX_OUT.relative_to(ROOT)}")
    print(f"Wrote {MD_OUT.relative_to(ROOT)}")
