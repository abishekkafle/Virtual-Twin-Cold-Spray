"""Create a non-destructive WebXR/KG interface add-on package for the AEI manuscript.

The script does not modify the latest manuscript. It creates a new folder with:
  * a Word addendum containing ready-to-paste interface-focused manuscript text;
  * primary WebXR interface screenshots;
  * candidate KG-linked photos/screenshots;
  * source/context files and an asset manifest.
"""

from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Tuple

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ABAQUS_ROOT = ROOT.parents[2]
KG_WEBXR_ROOT = ABAQUS_ROOT / "kg_driven_cold_spray"
SOURCE_MANUSCRIPT = (
    Path.home() / "Downloads" / "AEI_Reviewer_Revised_Cold_Spray_Virtual_Twin_Manuscript.docx"
)

OUT_STEM = "webxr_interface_addon_AEI"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_YELLOW = "FFF7D6"
WHITE = "FFFFFF"


PRIMARY_INTERFACE_FIGURES: List[Tuple[str, str, str]] = [
    (
        "paper_figs/fig5a_kg_browser.png",
        "Fig. A1a. Immersive research-KG browser",
        "The WebXR research panel presents curated material families and KG-resolved evidence fields inside the virtual cell.",
    ),
    (
        "paper_figs/fig5b_process_scene.png",
        "Fig. A1b. KG-bound process scene",
        "The A-Frame/Three.js cold-spray scene spatializes the gas supply, nozzle, particle stream, substrate and build-up while linking highlighted components to KG entities.",
    ),
    (
        "paper_figs/fig5c_knob_to_prediction.png",
        "Fig. A1c. Controller parameter to live prediction",
        "A pressure/temperature control interaction updates the surrogate readout through the browser-side model adapter.",
    ),
    (
        "paper_figs/fig5d_visual_response_LOW.png",
        "Fig. A1d. Low-response visual state",
        "The virtual substrate/coating state at a lower predicted deposition response, used as the baseline of the visual response comparison.",
    ),
    (
        "paper_figs/fig5d_visual_response_HIGH.png",
        "Fig. A1e. High-response visual state",
        "The virtual substrate/coating state at a higher predicted deposition response, showing the visual feedback path from prediction to scene.",
    ),
    (
        "paper_figs/fig5e_provenance_bridge.png",
        "Fig. A1f. Provenance bridge",
        "The interface exposes literature-backed provenance records, including source-paper identifiers and quoted context, next to the in-XR material view.",
    ),
    (
        "paper_figs/fig_webxr_evidence_round3.png",
        "Fig. A2. Evidence and provenance panel",
        "A higher-resolution interface capture showing how evidence rows and graph-derived context are surfaced to the operator-facing HMI.",
    ),
    (
        "paper_figs/fig_webxr_output_panel_round3.png",
        "Fig. A3. Runtime output panel",
        "The runtime output panel displays simulation-surrogate quantities and authorization status, separating numeric prediction from physical validation claims.",
    ),
    (
        "paper_figs/fig_webxr_deployment_gate_round3.png",
        "Fig. A4. Deployment gate strip",
        "The interface-level gate strip communicates standard, caution and blocked states before unsupported results are displayed.",
    ),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def copy_file(src: Path, dst: Path) -> dict:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return {
        "source": str(src),
        "package_path": str(dst),
        "size_bytes": dst.stat().st_size,
        "sha256": sha256(dst),
    }


def iter_images(root: Path) -> Iterable[Path]:
    for suffix in ("*.png", "*.jpg", "*.jpeg", "*.tif", "*.tiff", "*.webp"):
        yield from sorted(root.rglob(suffix))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", *, style=None, bold=False, italic=False, color=None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, bold=bold, italic=italic, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_callout(doc: Document, title: str, body: str, fill=PALE_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, color=RGBColor(40, 40, 40))
    set_table_borders(table, color="E2C044", size="8")
    add_para(doc, "")


def add_table(doc: Document, headers: List[str], rows: List[List[str]], widths: List[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    set_table_borders(table)
    add_para(doc, "")


def add_image_with_caption(doc: Document, image_path: Path, caption_title: str, caption_body: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    aspect = height_px / width_px if width_px else 1
    width_inches = 3.25 if aspect > 1.25 else 6.25
    shape = run.add_picture(str(image_path), width=Inches(width_inches))
    shape._inline.docPr.set("descr", caption_body)
    shape._inline.docPr.set("title", caption_title)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(10)
    r1 = cap.add_run(caption_title + ". ")
    set_run_font(r1, bold=True, size=9.5, color=NAVY)
    r2 = cap.add_run(caption_body)
    set_run_font(r2, size=9.5, italic=True, color=GRAY)


def make_docx(out_docx: Path, copied_primary: List[Tuple[Path, str, str]], manuscript_headings: List[str]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "AEI cold-spray virtual twin - WebXR interface add-on"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "Prepared as a non-destructive insertion package"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("WebXR Interface Addendum for the AEI Cold-Spray Virtual Twin Manuscript")
    set_run_font(r, size=22, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(
        "Ready-to-paste text, figure plan, captions and image assets for strengthening the KG-driven WebXR front-end contribution."
    )
    set_run_font(r, size=11.5, italic=True, color=GRAY)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Source manuscript", str(SOURCE_MANUSCRIPT)],
            ["Output type", "Addendum only; the attached manuscript is not edited."],
            ["Recommended manuscript action", "Add a front-end-centered subsection and a new multi-panel WebXR interface figure."],
            ["Claim boundary", "Interface demonstrates bounded model execution, authorization, evidence surfacing and KG-linked HMI context; no user-study or autonomous-control claims."],
        ],
        [1.75, 4.75],
    )

    add_callout(
        doc,
        "Editorial diagnosis",
        "The current manuscript is scientifically stronger than earlier versions, but the WebXR front end reads like an endpoint of the model pipeline rather than a primary contribution. For Advanced Engineering Informatics, the visible HMI, provenance-aware interaction design, and deployment gating can be the selling point if shown as an engineered interface layer rather than as decorative visualization.",
    )

    add_heading(doc, "1. Recommended insertion strategy", 1)
    add_para(
        doc,
        "Do not rewrite the whole paper. Add one focused interface subsection, strengthen the abstract/introduction language by one or two sentences, and insert one new multi-panel figure that shows the operator-facing WebXR twin. Keep the existing validation caution language intact.",
    )
    add_table(
        doc,
        ["Current location", "Recommended change", "Why it helps"],
        [
            [
                "Abstract",
                "Add one sentence naming the WebXR HMI as a deployable, provenance-aware interface rather than only a runtime wrapper.",
                "Signals the engineering-informatics contribution before reviewers reach the methods.",
            ],
            [
                "Section 3: System architecture",
                "Add a paragraph stating that the architecture terminates in three visible HMI surfaces: research-KG panel, spatial process scene and runtime authorization/model panel.",
                "Turns Figure 1 from a backend pipeline into an end-to-end virtual-twin system.",
            ],
            [
                "Section 6.3/6.4",
                "Insert a new subsection after browser replay verification: '6.4 WebXR front-end and KG-driven HMI'. Renumber the current traceability subsection if needed.",
                "Gives the interface enough manuscript real estate to become a paper contribution.",
            ],
            [
                "Results",
                "Add a short results paragraph describing what each interface surface demonstrates, without claiming user-performance benefits.",
                "Makes the photos/screenshots part of the evidence, not decoration.",
            ],
            [
                "Limitations",
                "Add one sentence that no formal user study or device-specific headset benchmark is included.",
                "Pre-empts an AEI reviewer asking for usability evidence.",
            ],
        ],
        [1.45, 2.75, 2.3],
    )

    add_heading(doc, "2. Ready-to-paste abstract addition", 1)
    add_para(
        doc,
        "Add after the sentence describing model authorization/runtime deployment:",
        italic=True,
        color=GRAY,
    )
    add_callout(
        doc,
        "Proposed abstract sentence",
        "The surrogate is deployed through a WebXR human-machine interface that couples a spatial cold-spray cell, a research-knowledge-graph browser, live model readouts and explicit standard/caution/blocked authorization states, thereby making model applicability and provenance visible at the point of interaction.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "3. Ready-to-paste introduction addition", 1)
    add_para(doc, "Add near the end of the Introduction, immediately before the contribution list:", italic=True, color=GRAY)
    add_callout(
        doc,
        "Proposed introduction paragraph",
        "For virtual twins in manufacturing, the scientific issue is not only whether a surrogate can reproduce a simulation dataset, but whether the resulting model can be exposed to users with traceable provenance and visible applicability limits. A browser- or headset-based interface is therefore not a cosmetic layer: it is where model scope, material identity, operator context and prediction uncertainty become operationally legible. In the present work, the WebXR front end is treated as part of the research contribution because it binds the KG, the qualified CEL dataset and the deployed surrogate into an inspectable interaction loop.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "4. Ready-to-paste Section 3 architecture addition", 1)
    add_callout(
        doc,
        "Proposed architecture paragraph",
        "The client layer is organized around three operator-visible surfaces. First, a research-KG panel exposes material records, evidence fields and source-paper context in the same environment in which predictions are requested. Second, a spatial process scene represents the gas supply, nozzle, particle stream, substrate and coating build-up as KG-bound components rather than as isolated graphics. Third, a model/authorization panel displays the simulation-surrogate outputs together with the runtime state that determines whether a query is standard, cautionary or blocked. This organization makes the WebXR front end a boundary-enforcing HMI rather than a passive visualization of numerical outputs.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "5. New subsection to add under Section 6", 1)
    add_heading(doc, "Suggested heading: 6.4 WebXR front-end and KG-driven HMI", 2)
    add_para(
        doc,
        "Insert this after the current browser replay verification subsection. If retained, the existing 'Traceability and HMI boundary' subsection can become 6.5.",
        italic=True,
        color=GRAY,
    )
    for paragraph in [
        "The WebXR layer operationalizes the virtual twin as an interactive HMI rather than as an offline prediction service. The implemented front end is a browser-delivered A-Frame/Three.js scene that can be viewed on a desktop browser or in a WebXR headset. It combines a spatial cold-spray cell with in-world panels for material/KG browsing, controller-driven process inputs, live surrogate readout and model-authorization state.",
        "The interface is KG-driven in two complementary senses. Material and operator context are loaded from graph-derived JSON/JSONL assets, so the research panel resolves materials, process attributes, procedures and evidence records from the graph rather than from hard-coded manuscript tables. At the same time, scene elements are named and highlighted through graph-facing component labels, allowing procedure steps and operator context to be displayed against the spatial cold-spray equipment model.",
        "The live model pathway exposes the simulation surrogate in the same environment. A parameter change updates the client-side model adapter, recomputes the surrogate outputs and refreshes the HMI panel. The interface deliberately separates prediction display from authorization display: standard in-domain queries are shown normally, caution-domain queries carry a constitutive-review warning, and unsupported material-pair or velocity requests are blocked. This prevents the WebXR view from visually normalizing predictions outside the evidence envelope.",
        "The WebXR implementation should be interpreted as a provenance-aware engineering-informatics interface, not as a validated operator-training study. No formal user experiment is reported. The contribution is the integrated design pattern: KG evidence, qualified simulation-surrogate inference and explicit deployment gates are brought into the same interactive scene, making the virtual twin inspectable at the point of use.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "6. Ready-to-paste results paragraph", 1)
    add_callout(
        doc,
        "Proposed results paragraph",
        "The WebXR captures in Figure 7 demonstrate the complete interface loop. The research-KG panel exposes material evidence and source context; the spatial process scene provides an equipment-level frame for the cold-spray interaction; parameter controls drive live surrogate updates; the scene responds visually to predicted changes; and the deployment gate communicates whether the request is standard, cautionary or blocked. These results verify interface integration and runtime behavior, not independent physical validation or user decision quality.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "7. Recommended new figure and captions", 1)
    add_para(
        doc,
        "Best option: add one new multi-panel figure after the current Figure 6 or replace Figure 6 with a richer interface figure. If the journal permits, call this Figure 7 to avoid disturbing the existing validation figures.",
    )
    add_table(
        doc,
        ["Panel", "Asset", "Manuscript role"],
        [
            ["a", "fig5a_kg_browser.png", "KG evidence/material browser"],
            ["b", "fig5b_process_scene.png", "Spatial cold-spray cell and component binding"],
            ["c", "fig5c_knob_to_prediction.png", "Controller input to model readout"],
            ["d", "fig5d_visual_response_LOW/HIGH.png", "Low/high visual response to predicted deposition state"],
            ["e", "fig5e_provenance_bridge.png", "Source-paper/provenance bridge"],
            ["f", "fig_webxr_deployment_gate_round3.png or fig_webxr_evidence_round3.png", "Authorization/evidence boundary shown at the interface"],
        ],
        [0.65, 2.55, 3.3],
    )
    add_callout(
        doc,
        "Proposed Figure 7 caption",
        "KG-driven WebXR interface for the cold-spray virtual twin. (a) The in-world research-KG browser presents material families and evidence fields derived from graph assets. (b) The spatial process scene represents the gas supply, nozzle, particle stream, substrate and coating build-up as KG-bound scene components. (c) Controller-driven process inputs update the browser-side surrogate readout. (d) Low- and high-response visual states show how predicted deposition response is communicated in the scene. (e) The provenance view links selected material context to source-paper evidence. (f) The deployment gate separates standard, caution and blocked queries before unsupported predictions are displayed. The interface verifies integrated model execution, provenance surfacing and authorization behavior; it does not constitute experimental validation or a user-study result.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "8. Suggested interface function table", 1)
    add_table(
        doc,
        ["Interface element", "Visible function", "Evidence/implementation boundary"],
        [
            [
                "Research-KG browser",
                "Surfaces material records, citations and evidence rows inside the WebXR scene.",
                "Supports provenance/context display; does not generate new simulation labels.",
            ],
            [
                "Spatial process scene",
                "Shows the cold-spray cell, component highlighting and deposition-state feedback.",
                "Supports interpretability and interaction; not a validated operator-training simulator.",
            ],
            [
                "Parameter controls",
                "Expose pressure/temperature-style controls and trigger live model updates.",
                "Pressure-to-gas-velocity coupling is provisional where noted; avoid physical calibration overclaims.",
            ],
            [
                "Model readout panel",
                "Displays simulation-surrogate outputs in the HMI.",
                "Outputs are simulation-surrogate quantities, not experimentally validated bond-quality metrics.",
            ],
            [
                "Authorization gate",
                "Shows standard/caution/blocked states and prevents unsupported predictions.",
                "Verifies software policy behavior; not a certification-grade safety system.",
            ],
        ],
        [1.5, 2.35, 2.65],
    )

    add_heading(doc, "9. Claim language to keep safe", 1)
    add_bullet(doc, "Use 'WebXR interface', 'front-end HMI', 'interactive virtual twin' and 'provenance-aware visualization'.")
    add_bullet(doc, "Use 'verified runtime replay' only for Python-to-JavaScript equivalence, not predictive accuracy.")
    add_bullet(doc, "Use 'authorization gate' or 'deployment boundary' rather than 'safety-certified controller'.")
    add_bullet(doc, "State that no formal user study, headset benchmark or experimental coating validation is included.")
    add_bullet(doc, "Treat OEM/operator manual photos as permission-check assets unless you own publication rights.")

    add_heading(doc, "10. Current manuscript anchor map", 1)
    add_para(doc, "The attached latest manuscript was inspected only to identify insertion anchors. It was not edited.")
    add_table(
        doc,
        ["Detected section anchor"],
        [[h] for h in manuscript_headings],
        [6.5],
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "11. Primary WebXR interface image gallery", 1)
    add_para(
        doc,
        "These are the primary assets recommended for manuscript use. They were copied into the package under `primary_webxr_interface_figures/`.",
    )
    for image_path, title, body in copied_primary:
        add_image_with_caption(doc, image_path, title, body)

    doc.save(out_docx)


def extract_manuscript_headings(path: Path) -> List[str]:
    if not path.exists():
        return ["Source manuscript not found during package build."]
    doc = Document(path)
    headings: List[str] = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            headings.append(text)
    return headings


def write_markdown(out_md: Path) -> None:
    out_md.write_text(
        """# WebXR interface manuscript addendum

This folder was created as a non-destructive add-on for the latest AEI manuscript. The source manuscript was not edited.

## Best manuscript move

Add a dedicated WebXR front-end/HMI subsection under Section 6 and add a new multi-panel interface figure. The interface can be a stronger selling point for Advanced Engineering Informatics if framed as a provenance-aware deployment and authorization layer, not as a cosmetic visualization.

## Primary assets

Use files under `primary_webxr_interface_figures/` for the main manuscript. The `candidate_kg_photos_permission_check/` folder contains KG-linked microstructure/operator images that may be useful for supplementary context but should be permission-checked before journal publication.

## Scientific boundary

The WebXR layer demonstrates integrated model execution, KG evidence surfacing, traceability and deployment-gate behavior. It does not demonstrate experimental validation, user-performance improvement, certification, or autonomous closed-loop control.
""",
        encoding="utf-8",
        newline="\n",
    )


def make_contact_sheet(copied_primary: List[Tuple[Path, str, str]], out_path: Path) -> None:
    thumb_w, thumb_h = 520, 325
    label_h = 58
    cols = 3
    rows = (len(copied_primary) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * thumb_w, rows * (thumb_h + label_h)), "white")
    draw = ImageDraw.Draw(sheet)
    try:
        font = ImageFont.truetype("arial.ttf", 18)
        small = ImageFont.truetype("arial.ttf", 15)
    except OSError:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    for idx, (image_path, title, _) in enumerate(copied_primary):
        col = idx % cols
        row = idx // cols
        x = col * thumb_w
        y = row * (thumb_h + label_h)
        with Image.open(image_path) as img:
            img = img.convert("RGB")
            img.thumbnail((thumb_w - 24, thumb_h - 24), Image.LANCZOS)
            px = x + (thumb_w - img.width) // 2
            py = y + (thumb_h - img.height) // 2
            sheet.paste(img, (px, py))
        draw.rectangle([x, y, x + thumb_w - 1, y + thumb_h + label_h - 1], outline=(210, 218, 226), width=2)
        draw.text((x + 12, y + thumb_h + 8), title.replace("Fig. ", ""), fill=(11, 37, 69), font=font)
        draw.text((x + 12, y + thumb_h + 34), image_path.name, fill=(90, 90, 90), font=small)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(out_path, quality=95)


def main() -> int:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_root = ROOT / f"{OUT_STEM}_{stamp}"
    out_root.mkdir(parents=True, exist_ok=False)

    copied: List[dict] = []
    missing: List[str] = []

    copied_primary: List[Tuple[Path, str, str]] = []
    for rel, title, body in PRIMARY_INTERFACE_FIGURES:
        src = KG_WEBXR_ROOT / rel
        dst = out_root / "primary_webxr_interface_figures" / Path(rel).name
        if src.exists():
            copied.append(copy_file(src, dst))
            copied_primary.append((dst, title, body))
        else:
            missing.append(str(src))

    for relroot, dest in [
        ("kg/MicroStructure", "candidate_kg_photos_permission_check/microstructure"),
        ("assets/troubleshooting", "candidate_kg_photos_permission_check/operator_machine"),
    ]:
        source_dir = KG_WEBXR_ROOT / relroot
        if not source_dir.exists():
            missing.append(str(source_dir))
            continue
        for src in iter_images(source_dir):
            dst = out_root / dest / src.relative_to(source_dir)
            copied.append(copy_file(src, dst))

    for rel in [
        "paper_figs/README.md",
        "paper_figs/captions.md",
        "paper_figs/SECTION7_SCOPE.md",
        "paper_figs/section7_supplement.md",
        "paper_figs/performance.md",
        "webxr_handover/README.md",
        "webxr_handover/webxr_integration_plan.md",
        "data/webxr_twin_data.json",
        "data/machine_kg.json",
        "cold_spray_cel_p5_model.js",
        "js/cs_model_adapter.js",
        "js/kg_client.js",
        "js/research_kg_panel.js",
        "js/materials_kg.js",
        "js/scene_adapter.js",
        "js/webxr_surrogate_runtime.js",
        "vr.html",
        "index.html",
    ]:
        src = KG_WEBXR_ROOT / rel
        dst = out_root / "supporting_webxr_context" / rel
        if src.exists():
            copied.append(copy_file(src, dst))
        else:
            missing.append(str(src))

    if SOURCE_MANUSCRIPT.exists():
        copied.append(
            copy_file(
                SOURCE_MANUSCRIPT,
                out_root / "source_latest_manuscript_DO_NOT_EDIT" / SOURCE_MANUSCRIPT.name,
            )
        )
    else:
        missing.append(str(SOURCE_MANUSCRIPT))

    readme = out_root / "README.md"
    readme.write_text(
        "\n".join(
            [
                "# WebXR interface add-on package for AEI manuscript",
                "",
                "This folder is intentionally separate from the manuscript. The attached manuscript was copied for context only and was not edited.",
                "",
                "## Contents",
                "",
                "- `AEI_WebXR_Interface_Addendum.docx`: ready-to-paste manuscript text, figure plan, captions and primary image gallery.",
                "- `primary_webxr_interface_figures/`: recommended WebXR/KG interface screenshots for the main paper.",
                "- `candidate_kg_photos_permission_check/`: KG-linked microstructure and operator/machine images. Use only after permission/copyright review.",
                "- `supporting_webxr_context/`: source files and notes that explain the WebXR implementation.",
                "- `source_latest_manuscript_DO_NOT_EDIT/`: copy of the latest manuscript for context only.",
                "",
                "## Recommended paper change",
                "",
                "Add a dedicated WebXR front-end/HMI subsection and one new multi-panel interface figure. Keep the existing validation limitations intact.",
                "",
            ]
        ),
        encoding="utf-8",
        newline="\n",
    )
    copied.append(
        {
            "source": "generated",
            "package_path": str(readme),
            "size_bytes": readme.stat().st_size,
            "sha256": sha256(readme),
        }
    )

    md = out_root / "AEI_WebXR_Interface_Addendum.md"
    write_markdown(md)
    copied.append(
        {
            "source": "generated",
            "package_path": str(md),
            "size_bytes": md.stat().st_size,
            "sha256": sha256(md),
        }
    )

    contact_sheet = out_root / "primary_webxr_interface_contact_sheet.png"
    make_contact_sheet(copied_primary, contact_sheet)
    copied.append(
        {
            "source": "generated",
            "package_path": str(contact_sheet),
            "size_bytes": contact_sheet.stat().st_size,
            "sha256": sha256(contact_sheet),
        }
    )

    docx = out_root / "AEI_WebXR_Interface_Addendum.docx"
    headings = extract_manuscript_headings(SOURCE_MANUSCRIPT)
    make_docx(docx, copied_primary, headings)
    copied.append(
        {
            "source": "generated",
            "package_path": str(docx),
            "size_bytes": docx.stat().st_size,
            "sha256": sha256(docx),
        }
    )

    manifest = {
        "schema_version": "1.0.0",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "package_root": str(out_root),
        "source_manuscript": str(SOURCE_MANUSCRIPT),
        "kg_webxr_root": str(KG_WEBXR_ROOT),
        "file_count": len(copied),
        "missing_sources": missing,
        "files": copied,
    }
    manifest_path = out_root / "ASSET_MANIFEST.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    zip_path = ROOT / f"{out_root.name}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for path in sorted(out_root.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(out_root.parent))

    print(f"Output folder: {out_root}")
    print(f"Output DOCX: {docx}")
    print(f"Output ZIP: {zip_path}")
    print(f"Copied/generated files: {len(copied)}")
    print(f"Missing sources: {len(missing)}")
    if missing:
        for item in missing:
            print(f"  {item}")
    return 0 if not missing else 2


if __name__ == "__main__":
    raise SystemExit(main())
