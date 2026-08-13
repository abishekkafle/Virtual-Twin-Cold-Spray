"""Create a non-destructive KG manuscript add-on package.

The package is designed for adding the missing knowledge-graph generation,
audit, ontology/provenance and KG-to-WebXR integration portions to the AEI
cold-spray virtual-twin manuscript.  It does not modify the manuscript.
"""

from __future__ import annotations

import csv
import hashlib
import json
import shutil
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ABAQUS_ROOT = ROOT.parents[2]
KG_ROOT = ABAQUS_ROOT / "cold_spray_kg - Copy"
KG_DATA = KG_ROOT / "data" / "ontology"
WEBXR_ROOT = ABAQUS_ROOT / "kg_driven_cold_spray"


def resolve_source_manuscript() -> Path:
    candidates = [
        Path.home() / "Downloads" / "AEI_Final_Application_Centered_Virtual_Twin_Manuscript_Minor_Cleanup.docx",
        Path.home() / "Downloads" / "AEI_Reviewer_Revised_Cold_Spray_Virtual_Twin_Manuscript.docx",
        ROOT / "paper" / "manuscript_nature_final.docx",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


SOURCE_MANUSCRIPT = resolve_source_manuscript()

OUT_STEM = "kg_manuscript_addon_AEI"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 90, 90)
BODY = RGBColor(30, 30, 30)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_YELLOW = "FFF7D6"


OPTIONAL_KG_FIGURES = [
    KG_ROOT / "docs" / "submission_aei" / "figures" / "Figure_1_claim_centered_representation.png",
    KG_ROOT / "docs" / "submission_aei" / "figures" / "Figure_2_decision_traces.png",
    KG_ROOT / "docs" / "submission_aei" / "figures" / "Figure_3_quantitative_evaluation.png",
    KG_ROOT / "docs" / "submission_aei" / "figures" / "Figure_4_evidence_readiness.png",
    KG_ROOT / "docs" / "submission_aei" / "figures" / "Figure_5_machine_aware_application.png",
    KG_ROOT / "Figures" / "Screenshot 2026-05-24 202526.png",
    KG_ROOT / "Figures" / "Screenshot 2026-05-24 202637.png",
    KG_ROOT / "Figures" / "Screenshot 2026-05-24 202804.png",
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def copy_file(src: Path, dst: Path) -> Dict[str, object]:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    return {
        "source": str(src),
        "package_path": str(dst),
        "size_bytes": dst.stat().st_size,
        "sha256": sha256(dst),
    }


def count_jsonl(path: Path) -> Dict[str, object]:
    stats: Dict[str, object] = {
        "file": str(path),
        "exists": path.exists(),
        "lines": 0,
        "bad_json": 0,
        "unique_source_paper": 0,
        "confidence_lt_1": 0,
        "top_types": [],
        "top_relations": [],
        "context_classes": [],
        "nli_models": [],
    }
    if not path.exists():
        return stats
    source_papers = set()
    types: Counter[str] = Counter()
    relations: Counter[str] = Counter()
    context_classes: Counter[str] = Counter()
    nli_models: Counter[str] = Counter()
    confidence_lt_1 = 0
    bad_json = 0
    lines = 0
    with path.open("r", encoding="utf-8", errors="replace") as stream:
        for line in stream:
            if not line.strip():
                continue
            lines += 1
            try:
                obj = json.loads(line)
            except json.JSONDecodeError:
                bad_json += 1
                continue
            if obj.get("source_paper"):
                source_papers.add(str(obj["source_paper"]))
            if obj.get("type"):
                types[str(obj["type"])] += 1
            if obj.get("relation"):
                relations[str(obj["relation"])] += 1
            if obj.get("context_class"):
                context_classes[str(obj["context_class"])] += 1
            if obj.get("nli_model"):
                nli_models[str(obj["nli_model"])] += 1
            if obj.get("confidence") is not None:
                try:
                    if float(obj["confidence"]) < 1.0:
                        confidence_lt_1 += 1
                except (TypeError, ValueError):
                    pass
    stats.update(
        {
            "lines": lines,
            "bad_json": bad_json,
            "unique_source_paper": len(source_papers),
            "confidence_lt_1": confidence_lt_1,
            "top_types": types.most_common(8),
            "top_relations": relations.most_common(8),
            "context_classes": context_classes.most_common(),
            "nli_models": nli_models.most_common(),
        }
    )
    return stats


def collect_kg_stats() -> Dict[str, object]:
    source_files = [
        "entities.jsonl",
        "entities.merged.jsonl",
        "triples.jsonl",
        "triples.audited.jsonl",
        "triples.nli_scored.jsonl",
        "operator_entities.jsonl",
        "operator_triples.jsonl",
        "bridge_triples.jsonl",
        "bridge_triples.prefilter.jsonl",
    ]
    webxr_files = [
        "entities.jsonl",
        "entities.merged.jsonl",
        "triples.jsonl",
        "triples.audited.jsonl",
        "operator_entities.jsonl",
        "operator_triples.jsonl",
        "bridge_triples.jsonl",
    ]
    return {
        "source_ontology": {name: count_jsonl(KG_DATA / name) for name in source_files},
        "webxr_kg_copy": {name: count_jsonl(WEBXR_ROOT / "kg" / name) for name in webxr_files},
    }


def write_csv(path: Path, headers: List[str], rows: List[List[object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(headers)
        writer.writerows(rows)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", *, style=None, bold=False, italic=False, color=None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic, color=color or BODY)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, color=BODY)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_callout(doc: Document, title: str, body: str, fill=PALE_YELLOW) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.45)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, color=BODY)
    set_table_borders(table, color="E2C044", size="8")
    add_para(doc, "")


def add_table(doc: Document, headers: List[str], rows: List[List[str]], widths: List[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(text))
            set_run_font(r, size=9.3, color=BODY)
    set_table_borders(table)
    add_para(doc, "")


def add_image_with_caption(doc: Document, image_path: Path, caption_title: str, caption_body: str, width=6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    shape._inline.docPr.set("descr", caption_body)
    shape._inline.docPr.set("title", caption_title)
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(10)
    r1 = cap.add_run(caption_title + ". ")
    set_run_font(r1, bold=True, size=9.5, color=NAVY)
    r2 = cap.add_run(caption_body)
    set_run_font(r2, italic=True, size=9.5, color=GRAY)


def make_pipeline_schematic(out_path: Path) -> None:
    width, height = 2400, 1500
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arial.ttf", 46)
        font = ImageFont.truetype("arial.ttf", 32)
        small = ImageFont.truetype("arial.ttf", 25)
    except OSError:
        title_font = ImageFont.load_default()
        font = ImageFont.load_default()
        small = ImageFont.load_default()
    draw.text((70, 55), "Knowledge-graph generation and binding path for the cold-spray virtual twin", fill=(11, 37, 69), font=title_font)
    boxes = [
        ("Literature corpus", "PDFs and paper-specific\nanalysis notes"),
        ("LLM-assisted extraction", "Claude Code drafts\npopulate_*.py scripts"),
        ("Ontology objects", "entities + triples with\nsource_paper and context"),
        ("Audit and scoring", "curator confidence +\nDeBERTa-v3 NLI fields"),
        ("Twin binding", "material crosswalk +\noperator bridge triples"),
        ("WebXR HMI", "KG evidence tier +\nmodel authorization display"),
    ]
    positions = [
        (110, 240),
        (890, 240),
        (1670, 240),
        (110, 620),
        (890, 620),
        (1670, 620),
    ]
    box_w, box_h = 610, 250
    for i, (head, body) in enumerate(boxes):
        x, y = positions[i]
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=28, fill=(232, 238, 245), outline=(46, 116, 181), width=5)
        draw.text((x + 30, y + 30), head, fill=(11, 37, 69), font=font)
        draw.multiline_text((x + 30, y + 98), body, fill=(45, 45, 45), font=small, spacing=10)
    # Arrows: top row left-to-right, then down and bottom row left-to-right.
    arrow_pairs = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5)]
    for start, end in arrow_pairs:
        sx, sy = positions[start]
        ex, ey = positions[end]
        if start == 2 and end == 3:
            x_mid = sx + box_w // 2
            y_start = sy + box_h + 12
            y_mid = ey - 45
            x_end = ex + box_w // 2
            draw.line((x_mid, y_start, x_mid, y_mid, x_end, y_mid, x_end, ey - 12), fill=(31, 77, 120), width=8)
            draw.polygon([(x_end - 18, ey - 12), (x_end + 18, ey - 12), (x_end, ey + 16)], fill=(31, 77, 120))
        else:
            ax = sx + box_w + 12
            ay = sy + box_h // 2
            bx = ex - 24
            by = ey + box_h // 2
            draw.line((ax, ay, bx, by), fill=(31, 77, 120), width=8)
            draw.polygon([(bx - 18, by - 17), (bx + 8, by), (bx - 18, by + 17)], fill=(31, 77, 120))
    callouts = [
        ("Preserved evidence boundary", "Prompt file is not preserved;\nreport surviving artifacts\nand the prompt caveat."),
        ("Do not mix counts", "129-paper raw/deployed exports\nand the 69-paper NLI subset\nare different KG stages."),
        ("Role in this paper", "KG provides provenance,\nmaterial identity, operator context\nand HMI evidence; not CEL labels."),
    ]
    y = 1010
    for i, (head, body) in enumerate(callouts):
        x = 120 + i * 740
        draw.rounded_rectangle((x, y, x + 660, y + 330), radius=20, fill=(255, 247, 214), outline=(226, 192, 68), width=4)
        draw.text((x + 28, y + 26), head, fill=(122, 90, 0), font=font)
        draw.multiline_text((x + 28, y + 92), body, fill=(45, 45, 45), font=small, spacing=8)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, quality=95)


def make_contact_sheet(images: List[Path], out_path: Path) -> None:
    existing = [p for p in images if p.exists()]
    if not existing:
        return
    thumb_w, thumb_h, label_h, cols = 620, 320, 76, 2
    rows = (len(existing) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * thumb_w, rows * (thumb_h + label_h)), "white")
    draw = ImageDraw.Draw(sheet)
    try:
        font = ImageFont.truetype("arial.ttf", 20)
        small = ImageFont.truetype("arial.ttf", 15)
    except OSError:
        font = ImageFont.load_default()
        small = ImageFont.load_default()
    for idx, path in enumerate(existing):
        col, row = idx % cols, idx // cols
        x, y = col * thumb_w, row * (thumb_h + label_h)
        with Image.open(path) as im:
            im = im.convert("RGB")
            im.thumbnail((thumb_w - 24, thumb_h - 24), Image.LANCZOS)
            px = x + (thumb_w - im.width) // 2
            py = y + (thumb_h - im.height) // 2
            sheet.paste(im, (px, py))
        draw.rectangle([x, y, x + thumb_w - 1, y + thumb_h + label_h - 1], outline=(210, 218, 226), width=2)
        draw.text((x + 12, y + thumb_h + 8), path.stem[:56], fill=(11, 37, 69), font=font)
        draw.text((x + 12, y + thumb_h + 38), path.name[:78], fill=(90, 90, 90), font=small)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(out_path, quality=95)


def extract_manuscript_headings(path: Path) -> List[str]:
    if not path.exists():
        return ["Source manuscript not found during package build."]
    doc = Document(path)
    headings = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if text and para.style and para.style.name.startswith("Heading"):
            headings.append(text)
    return headings


def read_json_manifest_summary() -> Dict[str, object]:
    manifest = ROOT / "webxr" / "cel_p5_kg_webxr_manifest.json"
    if not manifest.exists():
        return {}
    return json.loads(manifest.read_text(encoding="utf-8"))


def create_docx(
    out_docx: Path,
    stats: Dict[str, object],
    pipeline_img: Path,
    contact_sheet: Path,
    manuscript_headings: List[str],
    manifest: Dict[str, object],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "AEI cold-spray virtual twin - KG manuscript add-on"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.text = "Prepared as a non-destructive insertion package"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Knowledge-Graph Generation Addendum for the AEI Cold-Spray Virtual Twin Manuscript")
    set_run_font(r, size=22, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(
        "Ready-to-paste KG construction, ontology, audit, NLI-verification and KG-to-WebXR integration text."
    )
    set_run_font(r, size=11.5, italic=True, color=GRAY)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Source manuscript", str(SOURCE_MANUSCRIPT)],
            ["Output type", "Addendum only; the attached manuscript is not edited."],
            ["Recommended manuscript action", "Add a dedicated KG construction/audit subsection and a KG-to-WebXR binding paragraph."],
            [
                "Core claim boundary",
                "The KG is an evidence/provenance and HMI-context layer. It is not a source of new CEL labels, experimental validation or autonomous decisions.",
            ],
        ],
        [1.75, 4.75],
    )

    add_callout(
        doc,
        "Editorial diagnosis",
        "The current manuscript mentions the KG, but it does not yet explain how the graph was produced, what schema/provenance objects exist, how evidence support was audited, or how the graph is bound into the WebXR twin. For Advanced Engineering Informatics, that missing layer is a core contribution: it is the informatics bridge between literature evidence, simulation-surrogate scope and operator-facing HMI context.",
    )

    add_heading(doc, "1. Use the KG as a manuscript pillar", 1)
    add_para(
        doc,
        "Recommended framing: the paper has three coupled technical artifacts rather than two: a qualified CEL simulation-surrogate dataset, a bounded surrogate/runtime, and a provenance-aware KG layer that binds material identity, literature evidence and operator knowledge into the virtual twin."
    )
    add_bullet(doc, "Do not promote the KG as experimental validation.")
    add_bullet(doc, "Do not claim the KG adds new simulation labels.")
    add_bullet(doc, "Do claim the KG makes material identity, source evidence, uncertainty/status and operator context visible in the HMI.")
    add_bullet(doc, "Be explicit that KG artifact counts differ by stage: raw export, merged/audited subset, NLI-scored subset and WebXR-deployed copy.")

    add_heading(doc, "2. Evidence-supported KG facts", 1)
    source = stats["source_ontology"]
    webxr = stats["webxr_kg_copy"]
    rows = [
        [
            "Audited literature extraction subset",
            "extraction_protocol_findings.md",
            "69 papers; 1,894 entities; 3,631 triples; 64 notes; 70 populate scripts",
            "Use for the Methods description of the reconstructed extraction workflow.",
        ],
        [
            "Source raw literature ontology",
            "cold_spray_kg - Copy/data/ontology/entities.jsonl + triples.jsonl",
            f"{source['entities.jsonl']['lines']} entity lines; {source['triples.jsonl']['lines']} triple lines; {source['entities.jsonl']['unique_source_paper']} source_paper values in entities",
            "Use only if explicitly labelled as the raw/source ontology artifact.",
        ],
        [
            "Source NLI-scored audit subset",
            "triples.nli_scored.jsonl",
            f"{source['triples.nli_scored.jsonl']['lines']} triples; {source['triples.nli_scored.jsonl']['unique_source_paper']} papers; model: MoritzLaurer/DeBERTa-v3-large-mnli-fever-anli-ling-wanli",
            "Use to explain evidence-support scoring without inventing a threshold.",
        ],
        [
            "Operator-manual KG",
            "operator_entities.jsonl + operator_triples.jsonl",
            f"{source['operator_entities.jsonl']['lines']} operator entities; {source['operator_triples.jsonl']['lines']} operator triples",
            "Use for machine/HMI context, procedure steps, hazards and component relations.",
        ],
        [
            "Bridge layer",
            "bridge_triples.jsonl",
            f"{source['bridge_triples.jsonl']['lines']} corresponds_to bridge triples",
            "Use to explain how literature entities are linked to operator or machine-facing entities.",
        ],
        [
            "WebXR-deployed broad KG copy",
            "kg_driven_cold_spray/kg/entities.jsonl + triples.jsonl",
            f"{webxr['entities.jsonl']['lines']} entity lines; {webxr['triples.jsonl']['lines']} triple lines; {webxr['entities.jsonl']['unique_source_paper']} source_paper values",
            "Use to explain what the client-side WebXR materials/KG panels consume.",
        ],
    ]
    add_table(doc, ["KG layer", "Artifact", "Local count", "Manuscript usage"], rows, [1.45, 1.85, 1.65, 1.55])

    add_heading(doc, "3. Recommended insertion map", 1)
    add_table(
        doc,
        ["Manuscript location", "Add this KG content", "Reviewer value"],
        [
            [
                "Abstract",
                "One sentence stating that the WebXR twin is KG-grounded and provenance-aware.",
                "Signals AEI relevance before the paper becomes only CEL + ML.",
            ],
            [
                "Introduction",
                "Add KG as one of the main pillars, with explicit claim boundary.",
                "Prevents reviewers from seeing the KG as decoration.",
            ],
            [
                "Related work/background",
                "Add a paragraph on KG/provenance in digital twins and why source-aware evidence matters for manufacturing informatics.",
                "Places the contribution in AEI's digital-twin/provenance lane.",
            ],
            [
                "Methods",
                "Add a full subsection: KG extraction workflow, ontology artifacts, provenance fields, NLI scoring and human-in-the-loop caveats.",
                "Makes the graph reproducible and auditable.",
            ],
            [
                "System architecture / WebXR integration",
                "Add KG-to-surrogate binding through material crosswalk, bridge triples and HMI tiers.",
                "Connects graph construction to the virtual twin rather than leaving it as a side project.",
            ],
            [
                "Limitations",
                "Add caveats: prompt file not preserved; counts are stage-specific; KG evidence is provenance/context only.",
                "Pre-empts the most likely reviewer objections.",
            ],
        ],
        [1.35, 3.0, 2.15],
    )

    add_heading(doc, "4. Ready-to-paste abstract addition", 1)
    add_callout(
        doc,
        "Proposed abstract sentence",
        "The virtual twin is grounded by a cold-spray knowledge graph that records material entities, literature-derived process/property relations, source-paper context, NLI evidence-support scores and operator-machine bridge relations, allowing the WebXR HMI to expose provenance and applicability boundaries alongside simulation-surrogate outputs.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "5. Ready-to-paste introduction addition", 1)
    add_callout(
        doc,
        "Proposed introduction paragraph",
        "A central challenge for simulation-surrogate virtual twins is that numerical predictions are often separated from the evidence base that defines material identity, process context and applicability limits. In cold spray, this separation is especially problematic because material grade, powder morphology, gas condition, substrate state, heat treatment and measurement method vary across the literature. We therefore treat the knowledge graph as a first-class component of the virtual twin: it does not replace CEL simulations or supply new training labels, but it provides the provenance layer through which literature evidence, material crosswalks, operator-machine context and HMI warnings are made inspectable.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "6. New Methods subsection: KG construction and audit", 1)
    add_heading(doc, "Suggested heading: Knowledge-graph construction, provenance ontology and evidence audit", 2)
    for paragraph in [
        "The cold-spray KG was assembled as a line-delimited JSON knowledge graph rather than as an unstructured bibliography. The surviving extraction audit reconstructs a paper-level workflow in which each source paper was represented by an analysis note and a Python populate script. The populate scripts emit typed entities and typed relations using add_entity() and add_triple() calls, producing JSONL nodes and edges that can be loaded by the graph engine and by the WebXR client.",
        "The extraction protocol audit reports an audited literature subset of 69 papers, 1,894 entities and 3,631 triples. Across the surviving analysis notes, the schema version recorded at extraction evolves from v0.7.0 to v0.12.0, indicating a single developing ontology rather than disconnected ad hoc tables. The recovered workflow attributes the early cohort to claude-sonnet-4-6 and subsequent script authoring to claude-opus-4-7 in Claude Code sessions. Because the exact extraction prompt is not preserved as a standalone artifact, the manuscript should report the surviving evidence products--analysis notes, populate scripts, JSONL graph files and audit reports--rather than claiming full prompt-level reproducibility.",
        "Each graph edge stores the relation, head and tail identifiers/names, source_paper, context, confidence and curator_confidence fields. The context field carries the evidence sentence or extracted context used to justify the relation. This design makes provenance an edge-level property: a material, process parameter or property value can be traced back to the source-paper context that supported the assertion.",
        "Evidence support was further audited by an NLI-scored triple file. The triples.nli_scored.jsonl artifact contains 3,631 literature triples over 69 source papers and records DeBERTa-v3 MNLI/FEVER/ANLI/WANLI cross-encoder outputs, including confidence_nli, nli_dist and nli_hypothesis fields. These fields should be described as evidence-support scores and diagnostics, not as a guarantee of physical truth. The preserved package does not justify inventing a hard publication threshold if one is not explicitly recorded.",
        "The KG is connected to the virtual twin through two mechanisms. First, the material_kg_crosswalk maps simulation material identifiers to canonical KG material names with mapping_type and material_scope fields. Second, bridge_triples.jsonl provides corresponds_to relations between literature entities and operator or machine-facing entities. Together these layers allow the HMI to display material evidence, operator context and authorization status while preserving the distinction between exact canonical matches, grade/condition caveats and broader material-class mappings.",
    ]:
        add_para(doc, paragraph)

    add_heading(doc, "7. Ready-to-paste KG results paragraph", 1)
    add_callout(
        doc,
        "Proposed results paragraph",
        "The KG artifacts provide an auditable evidence layer for the virtual twin. The source ontology contains raw and merged literature entities/triples, an NLI-scored 69-paper audit subset, operator-manual entities/triples and 117 bridge relations. The deployed WebXR KG copy exposes 3,717 entity lines and 7,522 triple lines with source_paper provenance. In the manuscript, these counts should be reported as artifact-specific counts rather than collapsed into a single graph size because the raw, merged, audited and deployed files serve different roles in the pipeline.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "8. Ready-to-paste KG-to-WebXR integration paragraph", 1)
    add_callout(
        doc,
        "Proposed integration paragraph",
        "At runtime, the KG layer acts as the evidence tier of the WebXR twin. The client can load graph-derived material and operator assets, resolve a simulation material to a canonical KG material through the crosswalk, display source-paper context for material evidence and use bridge relations to connect literature entities to machine-facing components or procedures. This makes the interface provenance-aware: the user sees not only a simulation-surrogate response, but also the evidence tier and authorization boundary that constrain how the response should be interpreted.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "9. Recommended KG figure and caption", 1)
    add_image_with_caption(
        doc,
        pipeline_img,
        "Proposed KG Figure",
        "Knowledge-graph generation and binding path for the cold-spray virtual twin, from source papers and LLM-assisted populate scripts to JSONL entities/triples, NLI evidence-support scoring, material/operator bridge relations and WebXR HMI deployment.",
    )
    add_callout(
        doc,
        "Proposed caption text",
        "Knowledge-graph construction and virtual-twin binding workflow. Source PDFs were converted into paper-specific analysis notes and populate scripts; the scripts generated typed entities and triples with source_paper and context fields; an NLI-scored audit file added evidence-support diagnostics; material crosswalk and bridge triples then bound the literature/operator KG to the WebXR interface. Stage-specific counts are reported separately because raw, merged, audited and deployed KG artifacts serve different purposes.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "10. Suggested KG table caption", 1)
    add_callout(
        doc,
        "Proposed table caption",
        "Knowledge-graph artifact inventory and claim boundaries. Counts are line counts from the supplied JSONL artifacts and should not be interchanged across raw, merged, audited, NLI-scored and WebXR-deployed graph stages.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "11. Language to avoid", 1)
    add_bullet(doc, "Avoid: 'the KG validates the surrogate'. Better: 'the KG exposes provenance and applicability context for surrogate outputs'.")
    add_bullet(doc, "Avoid: '129 papers were fully NLI-audited' unless using the broader raw/deployed export. Better: 'the NLI-scored audit subset contains 3,631 triples over 69 papers'.")
    add_bullet(doc, "Avoid: 'LLM extraction is fully reproducible from prompts'. Better: 'surviving artifacts reconstruct the workflow; the standalone prompt is not preserved'.")
    add_bullet(doc, "Avoid: 'all material mappings are exact'. Better: 'mapping_type and material_scope preserve exact, grade/condition and broader-class distinctions'.")

    add_heading(doc, "12. Optional KG figure assets", 1)
    if contact_sheet.exists():
        add_image_with_caption(
            doc,
            contact_sheet,
            "Optional KG figure contact sheet",
            "Candidate figures from the prior KG manuscript package. Treat these as optional supporting assets; the generated KG pipeline figure above is the safest insertion figure for the current virtual-twin manuscript.",
        )

    add_heading(doc, "13. Current manuscript anchor map", 1)
    add_para(doc, "The latest manuscript was inspected only to identify insertion anchors. It was not edited.")
    add_table(doc, ["Detected section anchor"], [[h] for h in manuscript_headings], [6.5])

    add_heading(doc, "14. CEL-P5 manifest facts that connect KG to the twin", 1)
    kg_counts = manifest.get("kg_layer", {}).get("counts", {}) if manifest else {}
    hmi_tiers = manifest.get("hmi_tiers", []) if manifest else []
    add_table(
        doc,
        ["Manifest field", "Value"],
        [
            ["literature_entities_merged", kg_counts.get("literature_entities_merged", "not found")],
            ["literature_triples_audited_file", kg_counts.get("literature_triples_audited_file", "not found")],
            ["literature_triples_nli_scored", kg_counts.get("literature_triples_nli_scored", "not found")],
            ["operator_entities", kg_counts.get("operator_entities", "not found")],
            ["operator_triples", kg_counts.get("operator_triples", "not found")],
            ["bridge_triples", kg_counts.get("bridge_triples", "not found")],
            ["KG role in twin", manifest.get("kg_layer", {}).get("role_in_twin", "not found") if manifest else "not found"],
        ],
        [2.4, 4.1],
    )
    if hmi_tiers:
        add_table(
            doc,
            ["Tier", "Name", "Content"],
            [[tier.get("tier"), tier.get("name"), tier.get("content")] for tier in hmi_tiers],
            [0.55, 1.55, 4.4],
        )

    doc.save(out_docx)


def write_markdown(out_md: Path, stats: Dict[str, object]) -> None:
    source = stats["source_ontology"]
    webxr = stats["webxr_kg_copy"]
    out_md.write_text(
        f"""# KG manuscript addendum for the AEI cold-spray virtual twin paper

This folder is a separate, non-destructive package. It does not edit the latest manuscript.

## Core message

The KG should be treated as a first-class engineering-informatics contribution: it supplies provenance, material identity, source-paper context, operator knowledge and WebXR HMI evidence. It does not supply new CEL labels or experimental validation.

## Key artifact-specific counts

- Extraction audit subset: 69 papers, 1,894 entities and 3,631 triples.
- Source raw ontology: {source['entities.jsonl']['lines']} entity lines and {source['triples.jsonl']['lines']} triple lines.
- Source NLI-scored subset: {source['triples.nli_scored.jsonl']['lines']} triples over {source['triples.nli_scored.jsonl']['unique_source_paper']} source papers.
- Source operator KG: {source['operator_entities.jsonl']['lines']} operator entities and {source['operator_triples.jsonl']['lines']} operator triples.
- Bridge layer: {source['bridge_triples.jsonl']['lines']} corresponds_to bridge triples.
- WebXR-deployed broad copy: {webxr['entities.jsonl']['lines']} entity lines and {webxr['triples.jsonl']['lines']} triple lines.

## Best paper insertion

Add a dedicated Methods subsection named `Knowledge-graph construction, provenance ontology and evidence audit`, then add one KG-to-WebXR paragraph in the system architecture/integration section. Use the generated `kg_pipeline_schematic.png` as the safest new KG figure.

## Critical caveat

Do not mix graph counts across stages. Raw, merged, audited, NLI-scored and WebXR-deployed JSONL files are related but not identical.
""",
        encoding="utf-8",
        newline="\n",
    )


def make_readme(out_path: Path) -> None:
    out_path.write_text(
        """# KG manuscript add-on package

This folder was created to help add the missing knowledge-graph generation and provenance portions to the AEI cold-spray virtual-twin manuscript. The source manuscript is copied only for context and was not edited.

## Start here

1. `AEI_KG_Manuscript_Addendum.docx`
2. `AEI_KG_Manuscript_Addendum.md`
3. `tables/kg_artifact_inventory.csv`
4. `figures/kg_pipeline_schematic.png`
5. `kg_source_docs/extraction_protocol_findings.md`
6. `kg_source_docs/ENGINE_OVERVIEW.md`

## Claim boundary

The KG is an evidence/provenance layer and an HMI context layer. It should not be described as experimental validation, autonomous control, or a source of new simulation labels.
""",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_root = ROOT / f"{OUT_STEM}_{stamp}"
    out_root.mkdir(parents=True, exist_ok=False)

    copied: List[Dict[str, object]] = []
    missing: List[str] = []

    stats = collect_kg_stats()
    stats_path = out_root / "tables" / "kg_counts_summary.json"
    stats_path.parent.mkdir(parents=True, exist_ok=True)
    stats_path.write_text(json.dumps(stats, indent=2), encoding="utf-8")
    copied.append({"source": "generated", "package_path": str(stats_path), "size_bytes": stats_path.stat().st_size, "sha256": sha256(stats_path)})

    source = stats["source_ontology"]
    webxr = stats["webxr_kg_copy"]
    inventory_rows = [
        ["source/raw/entities.jsonl", KG_DATA / "entities.jsonl", source["entities.jsonl"]["lines"], source["entities.jsonl"]["unique_source_paper"], "Raw/source literature entities"],
        ["source/raw/triples.jsonl", KG_DATA / "triples.jsonl", source["triples.jsonl"]["lines"], source["triples.jsonl"]["unique_source_paper"], "Raw/source literature triples"],
        ["source/merged/entities.merged.jsonl", KG_DATA / "entities.merged.jsonl", source["entities.merged.jsonl"]["lines"], source["entities.merged.jsonl"]["unique_source_paper"], "Merged literature entities"],
        ["source/audited/triples.audited.jsonl", KG_DATA / "triples.audited.jsonl", source["triples.audited.jsonl"]["lines"], source["triples.audited.jsonl"]["unique_source_paper"], "Audited literature triple file"],
        ["source/nli/triples.nli_scored.jsonl", KG_DATA / "triples.nli_scored.jsonl", source["triples.nli_scored.jsonl"]["lines"], source["triples.nli_scored.jsonl"]["unique_source_paper"], "NLI-scored audit subset"],
        ["source/operator/operator_entities.jsonl", KG_DATA / "operator_entities.jsonl", source["operator_entities.jsonl"]["lines"], source["operator_entities.jsonl"]["unique_source_paper"], "Operator KG nodes"],
        ["source/operator/operator_triples.jsonl", KG_DATA / "operator_triples.jsonl", source["operator_triples.jsonl"]["lines"], source["operator_triples.jsonl"]["unique_source_paper"], "Operator KG edges"],
        ["source/bridge/bridge_triples.jsonl", KG_DATA / "bridge_triples.jsonl", source["bridge_triples.jsonl"]["lines"], source["bridge_triples.jsonl"]["unique_source_paper"], "Literature-to-operator bridge edges"],
        ["webxr/entities.jsonl", WEBXR_ROOT / "kg" / "entities.jsonl", webxr["entities.jsonl"]["lines"], webxr["entities.jsonl"]["unique_source_paper"], "WebXR-deployed broad entity copy"],
        ["webxr/triples.jsonl", WEBXR_ROOT / "kg" / "triples.jsonl", webxr["triples.jsonl"]["lines"], webxr["triples.jsonl"]["unique_source_paper"], "WebXR-deployed broad triple copy"],
    ]
    inventory_csv = out_root / "tables" / "kg_artifact_inventory.csv"
    write_csv(inventory_csv, ["label", "source_path", "line_count", "unique_source_paper", "manuscript_use"], inventory_rows)
    copied.append({"source": "generated", "package_path": str(inventory_csv), "size_bytes": inventory_csv.stat().st_size, "sha256": sha256(inventory_csv)})

    pipeline_img = out_root / "figures" / "kg_pipeline_schematic.png"
    make_pipeline_schematic(pipeline_img)
    copied.append({"source": "generated", "package_path": str(pipeline_img), "size_bytes": pipeline_img.stat().st_size, "sha256": sha256(pipeline_img)})

    contact_sheet = out_root / "figures" / "optional_kg_figures_contact_sheet.png"
    make_contact_sheet(OPTIONAL_KG_FIGURES, contact_sheet)
    if contact_sheet.exists():
        copied.append({"source": "generated", "package_path": str(contact_sheet), "size_bytes": contact_sheet.stat().st_size, "sha256": sha256(contact_sheet)})

    for src in OPTIONAL_KG_FIGURES:
        if src.exists():
            copied.append(copy_file(src, out_root / "figures" / "optional_kg_figures" / src.name))
        else:
            missing.append(str(src))

    for rel in [
        "ENGINE_OVERVIEW.md",
        "extraction_protocol_findings.md",
        "README.md",
        "VR_INTEGRATION.md",
        "CODEX_HANDOFF.md",
        "rater_app_inputs.md",
        "er_sample_30_rater_user.jsonl",
        "er_sample_30_rater_human2_Santosh2.jsonl",
        "multitail_sample_20_rater_user.jsonl",
    ]:
        src = KG_ROOT / rel
        if src.exists():
            copied.append(copy_file(src, out_root / "kg_source_docs" / rel))
        else:
            missing.append(str(src))

    for rel in [
        "entities.jsonl",
        "entities.merged.jsonl",
        "triples.jsonl",
        "triples.audited.jsonl",
        "triples.nli_scored.jsonl",
        "operator_entities.jsonl",
        "operator_triples.jsonl",
        "bridge_triples.jsonl",
        "bridge_triples.prefilter.jsonl",
    ]:
        src = KG_DATA / rel
        if src.exists():
            copied.append(copy_file(src, out_root / "kg_jsonl_artifacts" / "source_ontology" / rel))
        else:
            missing.append(str(src))

    for rel in [
        "webxr_handover/README.md",
        "webxr_handover/webxr_integration_plan.md",
        "data/webxr_twin_data.json",
        "data/machine_kg.json",
        "kg/entities.jsonl",
        "kg/triples.jsonl",
        "kg/entities.merged.jsonl",
        "kg/triples.audited.jsonl",
        "kg/operator_entities.jsonl",
        "kg/operator_triples.jsonl",
        "kg/bridge_triples.jsonl",
        "js/kg_client.js",
        "js/materials_kg.js",
        "js/research_kg_panel.js",
        "js/kg_overlay.js",
    ]:
        src = WEBXR_ROOT / rel
        if src.exists():
            copied.append(copy_file(src, out_root / "webxr_kg_bridge" / rel))
        else:
            missing.append(str(src))

    for rel in [
        "config/material_kg_crosswalk.csv",
        "webxr/cel_p5_kg_webxr_manifest.json",
        "reports/cel_p5_kg_webxr_integration_plan.md",
        "reports/cel_p5_webxr_runtime_verification.md",
    ]:
        src = ROOT / rel
        if src.exists():
            copied.append(copy_file(src, out_root / "virtual_twin_kg_bridge" / rel))
        else:
            missing.append(str(src))

    if SOURCE_MANUSCRIPT.exists():
        copied.append(copy_file(SOURCE_MANUSCRIPT, out_root / "source_latest_manuscript_DO_NOT_EDIT" / SOURCE_MANUSCRIPT.name))
    else:
        missing.append(str(SOURCE_MANUSCRIPT))

    manifest_json = read_json_manifest_summary()
    docx = out_root / "AEI_KG_Manuscript_Addendum.docx"
    create_docx(
        docx,
        stats,
        pipeline_img,
        contact_sheet,
        extract_manuscript_headings(SOURCE_MANUSCRIPT),
        manifest_json,
    )
    copied.append({"source": "generated", "package_path": str(docx), "size_bytes": docx.stat().st_size, "sha256": sha256(docx)})

    md = out_root / "AEI_KG_Manuscript_Addendum.md"
    write_markdown(md, stats)
    copied.append({"source": "generated", "package_path": str(md), "size_bytes": md.stat().st_size, "sha256": sha256(md)})

    readme = out_root / "README.md"
    make_readme(readme)
    copied.append({"source": "generated", "package_path": str(readme), "size_bytes": readme.stat().st_size, "sha256": sha256(readme)})

    package_manifest = {
        "schema_version": "1.0.0",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "package_root": str(out_root),
        "source_manuscript": str(SOURCE_MANUSCRIPT),
        "kg_root": str(KG_ROOT),
        "webxr_root": str(WEBXR_ROOT),
        "file_count": len(copied),
        "missing_sources": missing,
        "files": copied,
    }
    manifest_path = out_root / "PACKAGE_MANIFEST.json"
    manifest_path.write_text(json.dumps(package_manifest, indent=2), encoding="utf-8")

    zip_path = ROOT / f"{out_root.name}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for path in sorted(out_root.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(out_root.parent))

    print(f"Output folder: {out_root}")
    print(f"Output DOCX: {docx}")
    print(f"Output ZIP: {zip_path}")
    print(f"Copied/generated files: {len(copied)}")
    print(f"Missing sources: {len(missing)}")
    if missing:
        for item in missing:
            print(f"  {item}")
    return 0 if not missing else 2


if __name__ == "__main__":
    raise SystemExit(main())
